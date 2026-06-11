# agent/modules/knowledge/parsers/docx.py
"""DOCX parser — python-docx based, with Heading style detection."""

from __future__ import annotations

from typing import Optional

from agent.modules.knowledge.schemas import NormalizedDocument


def parse(
    raw: bytes,
    *,
    title: str = "",
    author: str = "",
    source_type: str = "project_doc",
    scope: str = "workspace",
    language: str = "zh",
    metadata: Optional[dict] = None,
) -> NormalizedDocument:
    try:
        import docx  # type: ignore
    except ImportError as e:
        return NormalizedDocument(
            title=title, author=author, source_type=source_type,
            scope=scope, language=language, format="docx",
            normalized_markdown="",
            metadata=metadata or {},
            warnings=[f"docx_parser_unavailable: {e!r}"],
        )
    import io
    try:
        d = docx.Document(io.BytesIO(raw))
    except Exception as e:
        return NormalizedDocument(
            title=title, author=author, source_type=source_type,
            scope=scope, language=language, format="docx",
            normalized_markdown="",
            metadata=metadata or {},
            warnings=[f"docx_open_failed: {e!r}"],
        )
    out_lines = []
    if not title:
        cp = d.core_properties
        if cp and cp.title:
            title = str(cp.title)
        if (not author) and cp and cp.author:
            author = str(cp.author)
    for para in d.paragraphs:
        style_name = (para.style.name or "").lower() if para.style else ""
        text = para.text
        if not text.strip() and not out_lines:
            continue
        if "heading 1" in style_name or "标题 1" in style_name:
            out_lines.append("# " + text)
        elif "heading 2" in style_name or "标题 2" in style_name:
            out_lines.append("## " + text)
        elif "heading 3" in style_name or "标题 3" in style_name:
            out_lines.append("### " + text)
        elif "heading 4" in style_name or "标题 4" in style_name:
            out_lines.append("#### " + text)
        elif "heading 5" in style_name or "标题 5" in style_name:
            out_lines.append("##### " + text)
        elif "heading 6" in style_name or "标题 6" in style_name:
            out_lines.append("###### " + text)
        else:
            out_lines.append(text)
        out_lines.append("")
    for table in d.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                out_lines.append("| " + " | ".join(cells) + " |")
        out_lines.append("")
    md = "\n".join(out_lines).strip()
    metadata = dict(metadata or {})
    metadata.setdefault("format_hint", "docx")
    if author and "author" not in metadata:
        metadata["author"] = author
    return NormalizedDocument(
        title=title,
        author=author,
        source_type=source_type,
        scope=scope,
        language=language,
        format="docx",
        normalized_markdown=md,
        metadata=metadata,
        warnings=[],
    )
